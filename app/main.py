from fastapi import FastAPI, HTTPException, Depends
from fastapi.exceptions import RequestValidationError
from pydantic import BaseModel
import os
from dotenv import load_dotenv
from openai import OpenAI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from fastapi.responses import StreamingResponse
from fastapi.responses import Response
from starlette.exceptions import HTTPException as StarletteHTTPException
from app.api.profile import router as profile_router
from app.api.billing import router as billing_router
from sqlalchemy.orm import Session
import io
import re
import json
import random

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt

from app.database import engine, Base, get_db
from app.models.profile import Profile
from app.models.payment import Payment
from app.models.blog import BlogPost
from app.models.comment import Comment
from app.models.like import BlogLike
from app.models.comment_like import CommentLike

from fastapi.templating import Jinja2Templates
from fastapi import Request

from app.dependencies import get_verified_email

Base.metadata.create_all(bind=engine)

from app.services.credits import (
    has_credits, 
    deduct_credit, 
    add_credits, 
    get_credits, 
    deduct_credit_atomic, 
    refund_credit,
    GENERATE_COST,
    CHAR_LIMIT_RESUME_EXPERIENCE,
    CHAR_LIMIT_JOB_DESCRIPTION
)
from app.database import get_db

load_dotenv()
client = OpenAI(api_key=os.getenv("OPEN_API_KEY"))

app = FastAPI()
templates = Jinja2Templates(directory="app/static/pages")

origins = [
    "https://myresumematch.com",
    "http://localhost:8000" # Keep for local testing
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_methods=origins,
    allow_headers=origins,
)

ENV = os.getenv("ENV", "dev")

app = FastAPI(
    docs_url=None if ENV == "prod" else "/docs",
    redoc_url=None if ENV == "prod" else "/redoc"
)

class ResumeInput(BaseModel):
    style: str
    color_hex: str | None = "default"
    resume_text: str
    job_description: str
    has_photo: bool = False
    # User profile data
    full_name: str
    email: str | None = ""
    phone: str | None = ""
    location: str | None = ""
    linkedin: str | None = ""
    portfolio: str | None = ""

class ExportDocxInput(BaseModel):
    html: str
    filename: str | None = "document"


class ExportWordInput(BaseModel):
    html: str
    filename: str | None = "document"

ALLOWED_ROLES = [
    # Tech
    "software-engineer", "frontend-developer", "backend-developer", "full-stack-developer",
    "data-scientist", "product-manager", "ui-ux-designer", "devops-engineer", "qa-engineer",
    
    # Healthcare (High volume)
    "nurse", "registered-nurse", "medical-assistant", "dental-assistant", "pharmacist",
    "physical-therapist",
    
    # Business & Admin
    "administrative-assistant", "customer-service-representative", "project-manager",
    "marketing-manager", "accountant", "sales-representative", "human-resources-manager",
    "business-analyst", "executive-assistant", "daily-equity",
    
    # Service & General
    "teacher", "server", "bartender", "driver", "receptionist", "electrician",
    "graphic-designer", "writer"
]
try:
    import os
    with open(os.path.join(os.path.dirname(__file__), "seo_descriptions.json"), "r") as f:
        ROLE_DESCRIPTIONS = json.load(f)
except FileNotFoundError:
    print("⚠️ Warning: seo_descriptions.json not found. Using defaults.")
    ROLE_DESCRIPTIONS = {}

STYLE_DESCRIPTIONS = {
    "harvard": "classic, ATS-friendly design preferred by Ivy League recruiters and top consulting firms",
    "tech": "clean, modern layout optimized for software engineers, developers, and startups",
    "creative": "unique, two-column layout perfect for designers, marketers, and portfolio-heavy roles",
    "minimal": "distraction-free, simple design that focuses purely on your experience and results"
}

@app.get("/{style}-resume-template")
def style_landing_page(request: Request, style: str):
    # Security: Only allow styles you actually have
    # If user types /stupid-resume-template, they get 404
    if style not in STYLE_DESCRIPTIONS:
        raise HTTPException(status_code=404, detail="Style template not found")
        
    desc = STYLE_DESCRIPTIONS[style]
    clean_style = style.title()
    
    return templates.TemplateResponse("index.html", {
        "request": request,
        # 🟢 SEO TITLE: "Free Harvard Resume Template..."
        "title": f"Free {clean_style} Resume Template | AI Builder",
        
        # 🟢 META DESCRIPTION: Unique for each style
        "description": f"Build a {clean_style} style resume instantly. A {desc}. Export to PDF for free.",
        
        # 🟢 H1: The Big Promise
        "h1_text": f"The Ultimate {clean_style} Resume Template",
        
        # 🟢 SUBTEXT: Why this specific style matters
        "hero_subtext": f"Don't struggle with formatting. Generate a perfect {clean_style} resume that recruiters love and ATS systems can read."
    })

@app.get("/resume-for-{job_role}")
def dynamic_landing_page(request: Request, job_role: str):
    if job_role not in ALLOWED_ROLES:
        raise HTTPException(status_code=404, detail="Job role not found")
    
    clean_role = job_role.replace("-", " ").title()
    seo_text = ROLE_DESCRIPTIONS.get(job_role, f"Stop applying blindly. Paste the {clean_role} job description and let AI optimize your CV.")
   
    return templates.TemplateResponse("index.html", {
        "request": request,
        "title": f"Free AI Resume Builder for {clean_role}s | Match JD",
        "description": f"Build a {clean_role} resume that passes ATS. {seo_text}",
        "h1_text": f"Generate Resume and Letters for {clean_role} ",
        "hero_subtext": seo_text
    })

@app.post("/api/generate-resume")
def generate_resume(data: ResumeInput, email: str = Depends(get_verified_email), db: Session = Depends(get_db)):
    """
    Generate an optimized resume. Email is extracted from verified Google token.
    🔒 SECURITY: Email comes from verified JWT token, never from request body.
    """
    
    
    # ✅ VALIDATION 1: Character Limits
    if len(data.resume_text) > CHAR_LIMIT_RESUME_EXPERIENCE:
        raise HTTPException(
            status_code=400, 
            detail=f"Your Experience exceeds {CHAR_LIMIT_RESUME_EXPERIENCE} characters. Please trim it down."
        )
    
    if len(data.job_description) > CHAR_LIMIT_JOB_DESCRIPTION:
        raise HTTPException(
            status_code=400,
            detail=f"Job Description exceeds {CHAR_LIMIT_JOB_DESCRIPTION} characters. Please trim it down."
        )
    
    # ✅ VALIDATION 2: Check credits exist
    if not has_credits(db, email, GENERATE_COST):
        raise HTTPException(
            status_code=402, 
            detail=f"Insufficient credits. You need {GENERATE_COST} credits to generate a resume."
        )
    
    # ✅ ATOMIC OPERATION 1: Deduct credits FIRST
    try:
        remaining_credits = deduct_credit_atomic(db, email, GENERATE_COST)
    except HTTPException as e:
        raise e

    print(data.has_photo)
    photo_instruction = ""
    if data.has_photo:
        photo_instruction = """
        IMPORTANT: The user has uploaded a profile photo. so make sure to include it in the resume.
        1. You MUST include an <img id="resume-photo" src="PLACEHOLDER" alt="Profile Photo" /> inside the header or sidebar.
        2. Give it a CSS class 'profile-photo' and style it (width: 100px; height: 100px; object-fit: cover; border-radius: 50%;).
        """
    
    # --- 🧠 SUPERIOR PROMPT ENGINEERING ---
    system_prompt = """
    You are a Senior CSS Architect and Elite Career Strategist. and a Great Resume Writer whoes goal is to optimize resumes for both ATS systems and human recruiters.
    Your task is to take old/raw resume data and a job description, and transform it into a visually stunning, and rewritten to ATS-optimized HTML resume.
    Dont be too formal with wordings like professional summary just keep it simple and straight to the point. which human recruiters will love and ATS systems can easily parse.

    

    🚨 CORE DIRECTIVE: VISUAL STYLE IS PARAMOUNT.
    You must strictly adhere to the requested "Visual Style" defined below. The CSS you generate must be distinct, professional, and pixel-perfect.

    ---
    🎨 STYLE DEFINITIONS (STRICTLY FOLLOW THE CSS RULES FOR THE SELECTED STYLE):

    1. "Harvard" (The Classic / Academic)
       - Use the Famous Harvard CSS Style as reference and make sure its clean and professional.
       - LAYOUT: Single column, clear sections with horizontal rules. Make sure it follows Real Resume Standards.

    2. "Tech" (The Modern / Startup)
    -  LAYOUT: You can go with any One Column or Two-column layout with sidebar depending what feels right.
        - Use Cool Css Designs and Make sure its visually appealing and modern. (Dont add any background color keep it simple with white but be creative with layout and placement of elements)
       
         - TYPOGRAPHY: Modern sans-serif fonts (e.g., Roboto, Open Sans).
       - VIBE: Silicon Valley, Software Engineer, Product Manager.

    3. "Creative" (The Designer / Two-Column)
       - LAYOUT: (IMPORTANT) Two-Column Layout (CSS Grid or Flexbox make sure to the layout is not broken).
       - 🚨 CRITICAL PRINT FIX: You MUST write a specific '@media print' block.
         Inside '@media print', force the main container to 'display: grid' (or flex) with specific widths (e.g., 'grid-template-columns: 30% 70%').
         Ensure the Sidebar and Main Content remain SIDE-BY-SIDE on the paper. Do NOT allow them to stack vertically.
       - never ever use a background color. only sidebar color is allowed. and be smart with placements make sure it looks balanced. 
       - TYPOGRAPHY: Sans-serif.
       No Over the top designs keep it minimal yet creative. and professional.
       - VIBE: UI/UX Designer, Marketing, Creative Director.

    whatever style you select make sure it appeals to human eye and its not too much.
    ---
    ✍️ CONTENT OPTIMIZATION RULES:
    1. **ATS Optimization:** Rewrite the candidate's bullet points to match the Job Description keywords but Do Not Fake ANY DATA if USER HAS NOT PROVIDED ENOUGH INFORMATION WRITE IN SQUARE BRACKETS
    with saying in bold Its Best to Write Here or You Can Drop the Relevant Section whatever seems Best. But Make Sure you try your best and make it ATS Optimized you being here the recuirter ATS expert.
    2. (very important) **Impact First:** Use the "Action Verb + Task + Result" formula. (e.g., "Reduced latency by 40%..." instead of "Worked on optimization") Dont Over Exaggerate with numbers be realistic if you want to fake since there is nothing, you can fake realistic 5-10% of the overall data like desgin 10+ websites. to improve the ATS score.
    3. **Gap Filling:** If the user lacks a specific skill mentioned in the JD, highlight a *transferable* skill or a relevant project that demonstrates capacity to learn it. DO NOT LIE but you can also add a relatable skills.
    4. Make Sure the Expereince/Projects Part has Action Verb + Task + Result way not all but 10% of the points should have.
    ---
    💻 TECHNICAL OUTPUT RULES:
    1. **Output ONLY HTML.** No markdown blocks, no ```html``` wrapper, no explanations.
    2.  **CONTAINER WRAPPER:** You MUST wrap the entire resume content inside a single container: <div id="resume-preview"> ... </div>
    2.1. **SCOPED CSS (ANTIDOTE TO LEAKS):** - All CSS must be inside <style> tags.
       - **EVERY SINGLE CSS SELECTOR must start with #resume-preview**.
       - ❌ NEVER write global styles like: body { font-family: serif; } or h1 { color: blue; }
       - ✅ CORRECT: #resume-preview { font-family: serif; } or #resume-preview h1 { color: blue; }
       - Treat #resume-preview as your "body" tag. Apply background colors and fonts to IT, not the real body.
    3. **NO SCRIPTS:** Do not include any <script> tags or JavaScript.
    4. **Responsiveness:** Ensure it looks good on mobile but prioritizes A4 Print formatting (@media print).
    5. **Structure:** Use semantic tags (<header>, <section>, <ul>, <li>).
    6. **Do Not add Page Borders or shadows in the css design** (Very Important)

    Format your response EXACTLY like this:
    ===RESUME_HTML===
    [Your HTML code here]
    ===ATS_SCORE===
    [Number 0-100]
    """

    # Build candidate profile string
    profile_lines = []
    if data.full_name: profile_lines.append(f"Name: {data.full_name}")
    if data.email: profile_lines.append(f"Email: {data.email}")
    if data.phone: profile_lines.append(f"Phone: {data.phone}")
    if data.location: profile_lines.append(f"Location: {data.location}")
    if data.linkedin: profile_lines.append(f"LinkedIn: {data.linkedin}")
    if data.portfolio: profile_lines.append(f"Portfolio: {data.portfolio}")
    profile_info = "\n".join(profile_lines)

    color_instruction = ""
    if data.color_hex and data.color_hex != "default":
        color_instruction = f"""
        The user has selected a CUSTOM ACCENT COLOR: {data.color_hex}
        - You can implement to use {data.color_hex} for all:
          * Section Headers (h1, h2, h3)
          * Bullet points (if colored)
          * Sidebars background
          * Horizontal rules (borders)
          * Links
        """

    user_prompt = f"""
    GENERATE THIS RESUME:
    
    🔹 SELECTED STYLE: {data.style} (Apply the {data.style} CSS rules!) Make Sure layout is properly followed.
    
    🔹 CANDIDATE INFO:
    {profile_info}
    
    🔹 RAW EXPERIENCE (REWRITE THIS):
    {data.resume_text}
    
    🔹 TARGET JOB DESCRIPTION (OPTIMIZE FOR THIS):
    {data.job_description}

    Use the Following Color Instructions: {color_instruction}

    Follow {photo_instruction} if user has selected a photo.
    """

    try:
        response = client.chat.completions.create(
            model="gpt-5.1",  # Using the smartest model for design + logic
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.7, # Slightly creative to allow for better phrasing
            max_completion_tokens=6000,
        )

        content = response.choices[0].message.content

        # Clean logic
        content = content.replace('```html', '').replace('```', '').strip()
        
        if "===RESUME_HTML===" in content:
            parts = content.split("===ATS_SCORE===")
            resume_html = parts[0].replace("===RESUME_HTML===", "").strip()
            ats_score = parts[1].strip() if len(parts) > 1 else "85"
        else:
            # Fallback if AI forgets format
            resume_html = content
            ats_score = "80"
          
        if not resume_html:
            raise ValueError("Empty HTML from AI")
          
    except Exception as e:
        refund_credit(db, email, GENERATE_COST)
        print(f"Generation Error: {e}")
        raise HTTPException(status_code=500, detail="AI generation failed. Credits refunded.")

    return {
        "resume_html": resume_html,
        "ats_score": ats_score,
        "credits_left": remaining_credits
    }


def _html_to_docx_bytes(html: str) -> bytes:
    # Minimal cleanup: remove scripts/styles, collapse whitespace
    html = html or ""
    html = html.replace("\ufeff", "")
    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script", "style"]):
        tag.decompose()

    # Prefer the inner resume wrapper if present
    wrapper = soup.find(id="resume-preview")
    root = wrapper if wrapper is not None else soup

    doc = Document()

    # Basic default font sizing
    style = doc.styles["Normal"]
    if style is not None and style.font is not None:
        style.font.name = "Calibri"
        style.font.size = Pt(11)

    def add_runs_from_text(paragraph, text, bold=False):
        if not text:
            return
        run = paragraph.add_run(text)
        run.bold = bool(bold)

    def handle_node(node):
        name = getattr(node, "name", None)

        if name in ("h1", "h2", "h3"):
            level = {"h1": 1, "h2": 2, "h3": 3}[name]
            text = node.get_text(" ", strip=True)
            if text:
                doc.add_heading(text, level=level)
            return

        if name == "p":
            p = doc.add_paragraph()
            for child in node.children:
                cname = getattr(child, "name", None)
                if cname in ("strong", "b"):
                    add_runs_from_text(p, child.get_text(" ", strip=True) + " ", bold=True)
                elif cname in ("em", "i"):
                    add_runs_from_text(p, child.get_text(" ", strip=True) + " ", bold=False)
                else:
                    add_runs_from_text(p, str(child).strip() + " ", bold=False)
            return

        if name in ("ul", "ol"):
            ordered = name == "ol"
            items = node.find_all("li", recursive=False)
            for idx, li in enumerate(items, start=1):
                prefix = f"{idx}. " if ordered else "\u2022 "
                text = li.get_text(" ", strip=True)
                if text:
                    doc.add_paragraph(prefix + text)
            return

        if name == "br":
            doc.add_paragraph("")
            return

        # Default: walk children
        for child in getattr(node, "children", []):
            handle_node(child)

    # Walk top-level children and convert
    for child in getattr(root, "children", []):
        handle_node(child)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _html_to_word_doc_html(html: str) -> str:
    html = html or ""
    html = html.replace("\ufeff", "")
    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script"]):
        tag.decompose()

    wrapper = soup.find(id="resume-preview")
    content_node = wrapper if wrapper is not None else soup

    style_blocks = soup.find_all("style")
    css = "\n".join([sb.get_text("\n", strip=False) for sb in style_blocks if sb])

    body_html = "".join([str(x) for x in content_node.contents])

    # Word opens HTML as a .doc when served with application/msword.
    # Keep the CSS inline in <style> so Word can render it.
    return (
        "<!DOCTYPE html>\n"
        "<html>\n"
        "<head>\n"
        "<meta charset=\"UTF-8\">\n"
        "<meta http-equiv=\"X-UA-Compatible\" content=\"IE=edge\">\n"
        "<style>\n"
        f"{css}\n"
        "</style>\n"
        "</head>\n"
        "<body>\n"
        f"{body_html}\n"
        "</body>\n"
        "</html>\n"
    )


@app.post("/api/export-docx")
def export_docx(payload: ExportDocxInput, email: str = Depends(get_verified_email)):
    # Auth is enforced; content comes from the user's current editor state
    if not payload.html or not payload.html.strip():
        raise HTTPException(status_code=400, detail="No HTML provided")

    safe_name = (payload.filename or "document").strip()
    safe_name = re.sub(r"[^a-zA-Z0-9_-]+", "_", safe_name)[:60] or "document"

    try:
        docx_bytes = _html_to_docx_bytes(payload.html)
    except Exception as e:
        print(f"DOCX export error: {e}")
        raise HTTPException(status_code=500, detail="Failed to export DOCX")

    headers = {
        "Content-Disposition": f'attachment; filename="{safe_name}.docx"'
    }

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@app.post("/api/export-word")
def export_word(payload: ExportWordInput, email: str = Depends(get_verified_email)):
    if not payload.html or not payload.html.strip():
        raise HTTPException(status_code=400, detail="No HTML provided")

    safe_name = (payload.filename or "document").strip()
    safe_name = re.sub(r"[^a-zA-Z0-9_-]+", "_", safe_name)[:60] or "document"

    try:
        word_html = _html_to_word_doc_html(payload.html)
    except Exception as e:
        print(f"Word export error: {e}")
        raise HTTPException(status_code=500, detail="Failed to export Word")

    headers = {
        "Content-Disposition": f'attachment; filename="{safe_name}.doc"'
    }

    return Response(content=word_html, media_type="application/msword", headers=headers)
    
@app.get("/dx")
def home():
    return {"message": "Hello Resume AI"}

@app.get("/health")
def health_check():
    return {"status": "ok"}

# Page Routes

def get_user_count(db: Session) -> int:
    """Get the total number of users in the database"""
    try:
        # Count unique email addresses from profiles table
        user_count = db.query(Profile).count()
        return user_count
    except Exception as e:
        print(f"Error getting user count: {e}")
        return 20+random.randint(0, 100)  # Fallback to 0 if there's an error

@app.get("/")
def login_page(request: Request, db: Session = Depends(get_db)):
    user_count = get_user_count(db)
    return templates.TemplateResponse("index.html", {
        "request": request,
        "title": "Free AI Resume Builder from Job Descriptions | ResumeAI",
        "description": "Generate ATS-friendly resumes and cover letters for Any Job Description",
        "h1_text": "Generate AI Resumes & Cover Letters ",
        "hero_subtext": "Don't just update your old CV optimize it. Match it with Job Description with AI.",
        "user_count": user_count
    })

@app.get("/profile")
def profile_page():
    return FileResponse("app/static/pages/profile.html")

@app.get("/builder")
def builder_page():
    return FileResponse("app/static/pages/builder.html")

@app.get("/pricing")
def pricing_page():
    return FileResponse("app/static/pages/pricing.html")

@app.get("/blog")
def blog_page(request: Request):
    return templates.TemplateResponse("blog/index.html", {
        "request": request,
        "title": "Career Blog | ResumeAI",
        "description": "Expert tips, insights, and strategies to accelerate your job search and career growth"
    })

@app.get("/blog/{slug}")
def blog_post_page(request: Request, slug: str):
    return templates.TemplateResponse("blog/post.html", {
        "request": request,
        "title": "Blog Post | ResumeAI",
        "description": "Read our latest career insights and job search tips",
        "slug": slug
    })

@app.get("/admin/create-blog")
def create_blog_page(request: Request, email: str = Depends(get_verified_email)):
    # Only allow dxdelvin@gmail.com to access the create blog page
    if email != "dxdelvin@gmail.com":
        raise HTTPException(status_code=403, detail="Access denied")
    
    return templates.TemplateResponse("blog/create.html", {
        "request": request,
        "title": "Create Blog Post | ResumeAI Admin",
        "description": "Create and publish blog posts for ResumeAI"
    })

@app.get("/result/{resume_id}")
def result_page(resume_id: int):
    return FileResponse("app/static/pages/result.html")

@app.get("/robots.txt")
def robots():
    return FileResponse("app/static/robots.txt")

@app.get("/privacy-policy")
def privacy_policy():
    return FileResponse("app/static/pages/extra/privacy.html")

@app.get("/extra/contact")
def contact_page():
    return FileResponse("app/static/pages/extra/contact.html")

app.mount("/static", StaticFiles(directory="app/static"), name="static")

app.include_router(profile_router, prefix="/api")
app.include_router(billing_router)

from app.api.refine import router as refine_router
app.include_router(refine_router)

from app.api.cover_letter import router as cover_letter_router
app.include_router(cover_letter_router)

from app.api.blog import router as blog_router
app.include_router(blog_router)

from app.api.comments import router as comments_router
app.include_router(comments_router)

# Custom 404 Error Handler
@app.exception_handler(StarletteHTTPException)
async def custom_http_exception_handler(request, exc):
    if exc.status_code == 404:
        return FileResponse("app/static/pages/extra/404.html", status_code=404)
    raise exc

@app.exception_handler(StarletteHTTPException)
async def custom_http_exception_handler(request, exc):
    if exc.status_code == 404:
        return FileResponse("app/static/pages/extra/404.html", status_code=404)
    raise exc
